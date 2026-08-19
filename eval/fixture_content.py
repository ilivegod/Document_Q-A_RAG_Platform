"""Deterministic agency brief content for eval fixtures."""

from pathlib import Path

from docx import Document as DocxDocument

BRIEF_SECTIONS: list[tuple[str, list[str]]] = [
    (
        "Client Overview",
        [
            "Client name: Acme Corp",
            "Engagement: EHR MVP for a regional clinic network",
            "Primary users: clinicians and front-desk staff",
        ],
    ),
    (
        "Commercial Terms",
        [
            "Liability cap: $50,000 for direct damages",
            "Deposit: 30% due upon SOW acceptance",
            "Agency hourly rate for change orders: $150 per hour",
            "Target delivery timeline: 12 weeks from kickoff",
        ],
    ),
    (
        "Technology Stack",
        [
            "Frontend: React with TypeScript",
            "Backend API: FastAPI on Python 3.12",
            "Database: PostgreSQL 16 with pgvector for search",
            "Hosting: Fly.io for API, Vercel for frontend",
        ],
    ),
    (
        "In Scope",
        [
            "Patient search and chart summary view",
            "Appointment scheduling for clinic staff",
            "Role-based access for admin, clinician, and reception roles",
            "Export patient demographics to CSV",
        ],
    ),
    (
        "Out of Scope",
        [
            "Native iOS or Android mobile applications",
            "HIPAA compliance audit or certification services",
            "Third-party billing system integrations",
            "Post-handoff maintenance retainers",
        ],
    ),
]


def write_agency_sample_docx(path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = DocxDocument()
    doc.add_heading("Agency Sample Brief — Acme Corp EHR MVP", level=0)
    for heading, bullets in BRIEF_SECTIONS:
        doc.add_heading(heading, level=1)
        for line in bullets:
            doc.add_paragraph(line, style="List Bullet")
    doc.save(path)
